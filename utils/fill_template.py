
import os
from docx import Document
from datetime import datetime
from docx2pdf import convert
import json

def fill_docx(template_path, context, output_path):
    doc = Document(template_path)
    for p in doc.paragraphs:
        for key, value in context.items():
            p.text = p.text.replace(f"{{{{{key}}}}}", str(value or " "))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value or " "))
    doc.save(output_path)

def get_next_contract_number():
    path = "last_number.txt"
    if not os.path.exists(path): return "001"
    with open(path, "r+") as f:
        number = int(f.read().strip()) + 1
        f.seek(0); f.write(str(number)); f.truncate()
    return str(number).zfill(3)

def calculate_price(service_name, kwt):
    with open("pricing.json", "r", encoding="utf-8") as f:
        prices = json.load(f)
    service = prices.get(service_name, {})
    if service.get("type") == "per_kwt":
        try: return round(float(kwt.replace(",", ".")) * service["price_per_kwt"], 2)
        except: return 0
    return service.get("price", 0)

def generate_contract_pdf(data):
    os.makedirs("generated", exist_ok=True)
    number = get_next_contract_number()
    date = datetime.now().strftime("%d.%m.%Y")
    price = calculate_price(data["УСЛУГА"], data.get("МОЩНОСТЬ", ""))
    context = {
        "номер": number,
        "ДАТА": date,
        "ФИО": data.get("ФИО", ""),
        "ПАСПОРТ": data.get("ПАСПОРТ", ""),
        "ИНН": data.get("ИНН", ""),
        "АДРЕС": data.get("АДРЕС", ""),
        "ТЕЛЕФОН": data.get("ТЕЛЕФОН", ""),
        "EMAIL": data.get("EMAIL", ""),
        "УСЛУГА": data.get("УСЛУГА", ""),
        "МОЩНОСТЬ": data.get("МОЩНОСТЬ", ""),
        "СТОИМОСТЬ_УСЛУГИ": f"{price:,.2f}".replace(",", " ").replace(".", ",")
    }
    docx_file = f"generated/dogovor_{number}.docx"
    pdf_file = f"generated/dogovor_{number}.pdf"
    fill_docx("templates/contract_template.docx", context, docx_file)
    convert(docx_file, pdf_file)
    return pdf_file
